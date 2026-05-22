"""Multi-format document parser with caching and error isolation."""

import subprocess
import tempfile
from pathlib import Path

from rag_system.config import CACHE_DIR
from rag_system.utils import (
    file_hash,
    is_cached,
    logger,
    read_cache,
    safe_text,
    write_cache,
)


def parse_file(filepath: Path) -> str | None:
    """Parse a single file to plaintext. Returns None on failure.
    Uses filesystem cache keyed by file content hash for idempotent ingest.
    """
    if is_cached(filepath, CACHE_DIR):
        logger.debug("Cache hit: %s", filepath.name)
        return read_cache(filepath, CACHE_DIR)

    ext = filepath.suffix.lower()
    text: str | None = None

    try:
        if ext in (".docx", ".docm"):
            text = _parse_docx(filepath)
        elif ext == ".doc":
            text = _parse_doc(filepath)
        elif ext == ".xlsx":
            text = _parse_xlsx(filepath)
        elif ext == ".xls":
            text = _parse_xls(filepath)
        elif ext == ".pdf":
            text = _parse_pdf(filepath)
        else:
            logger.warning("Unsupported extension: %s", ext)
            return None
    except Exception:
        logger.exception("Failed to parse: %s", filepath.name)
        return None

    if not text or not text.strip():
        logger.warning("Empty content: %s", filepath.name)
        return None

    text = safe_text(text)
    write_cache(filepath, CACHE_DIR, text)
    return text


def _parse_docx(filepath: Path) -> str:
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            paragraphs.append(t)

    # Also check tables (some scripts have embedded tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _parse_doc(filepath: Path) -> str | None:
    """Try python-docx first (some .doc are actually .docx), then libreoffice fallback."""
    # Many .doc files are secretly .docx (OOXML)
    try:
        return _parse_docx(filepath)
    except Exception:
        pass

    # Fall back to LibreOffice headless conversion
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmpdir,
                    str(filepath),
                ],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                logger.warning(
                    "LibreOffice conversion failed for %s: %s",
                    filepath.name, result.stderr.strip(),
                )
                return None

            converted = Path(tmpdir) / f"{filepath.stem}.docx"
            if converted.exists():
                return _parse_docx(converted)
    except FileNotFoundError:
        logger.warning(
            "LibreOffice not installed, skipping .doc file: %s", filepath.name,
        )
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice timed out on: %s", filepath.name)

    return None


def _parse_xlsx(filepath: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(filepath, read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"【工作表：{sheet_name}】")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _parse_xls(filepath: Path) -> str:
    import xlrd

    wb = xlrd.open_workbook(str(filepath))
    lines = []
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        lines.append(f"【工作表：{sheet_name}】")
        for row_idx in range(ws.nrows):
            cells = [
                str(ws.cell_value(row_idx, col_idx)).strip()
                for col_idx in range(ws.ncols)
                if str(ws.cell_value(row_idx, col_idx)).strip()
            ]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _parse_pdf(filepath: Path) -> str:
    import pdfplumber

    with pdfplumber.open(str(filepath)) as pdf:
        pages = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
