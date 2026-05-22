"""Format generated script into a structured .docx file for reading aloud.

Design System: Swiss Editorial + Noto Sans SC typography
  - Font: 等线 (DengXian) — Windows 10+ default, cleaner than 微软雅黑
  - Body: 12pt / 1.8x line height — comfortable reading distance
  - High contrast: #1E293B text on white
  - One sentence per line — each line = one breath unit
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_LINE_SPACING


# ═══ Design System: Swiss Editorial ═══

FONT_FAMILY = "等线"          # DengXian — clean modern sans, Win10+ default
FONT_FALLBACK = "微软雅黑"    # Fallback for older Windows

# Size hierarchy
SIZE_COVER = Pt(16)           # Cover title
SIZE_BODY = Pt(12)            # Voiceover body — comfortable reading aloud
SIZE_BRIEF = Pt(11)           # Summary
SIZE_HUAZI = Pt(10)           # On-screen text annotations
SIZE_HASHTAG = Pt(9)          # Social tags

# Color palette
COLOR_TEXT = RGBColor(0x1E, 0x29, 0x3B)      # Slate-800 — softer than pure black
COLOR_TEXT_BOLD = RGBColor(0x0F, 0x17, 0x2A)  # Slate-900 — cover title
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)      # Slate-500 — brief
COLOR_HUAZI = RGBColor(0xC2, 0x41, 0x0C)      # Orange-700 — key specs
COLOR_HASHTAG = RGBColor(0x1D, 0x4E, 0xD8)    # Blue-700 — tags

# Spacing (in Pt)
SPACE_AFTER_COVER = Pt(14)
SPACE_AFTER_BRIEF = Pt(10)
SPACE_AFTER_BODY = Pt(5)      # Breath between spoken lines
SPACE_AFTER_HUAZI = Pt(2)
LINE_SPACING_BODY = 1.8       # Comfortable reading aloud distance


def _apply_font(run, size: Pt, bold: bool = False, color: RGBColor = COLOR_TEXT):
    """Apply font styling to a run with fallback support."""
    run.font.name = FONT_FAMILY
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    # East-Asian font fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', FONT_FAMILY)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', FONT_FAMILY)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', FONT_FAMILY)


def format_script_to_docx(
    text: str,
    product_name: str,
    persona: str,
    key_points: str,
    output_path: Path,
) -> Path:
    parsed = _parse_llm_output(text, product_name, persona)

    doc = Document()

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = SIZE_BODY
    style.font.color.rgb = COLOR_TEXT
    pf = style.paragraph_format
    pf.space_after = SPACE_AFTER_BODY
    pf.line_spacing = LINE_SPACING_BODY

    # ── 封面 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = SPACE_AFTER_COVER
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(f"封面：{parsed['cover']}")
    _apply_font(run, SIZE_COVER, bold=True, color=COLOR_TEXT_BOLD)

    # ── 简介 ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = SPACE_AFTER_BRIEF
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(f"简介：{parsed['brief']}")
    _apply_font(run, SIZE_BRIEF, bold=False, color=COLOR_MUTED)

    # ── Body — post-process: strip prices from voiceover, merge 花字 ──
    merged_body = _merge_standalone_huazi(parsed["body"])
    merged_body = _strip_prices_from_body(merged_body)

    for line in merged_body:
        if not line.strip():
            continue

        main_text, huazi_note = _split_huazi_from_line(line.strip())

        p = doc.add_paragraph()
        p.paragraph_format.space_after = SPACE_AFTER_BODY
        p.paragraph_format.line_spacing = LINE_SPACING_BODY

        # Main voiceover text
        if main_text:
            run = p.add_run(main_text)
            _apply_font(run, SIZE_BODY, bold=False, color=COLOR_TEXT)

        # Inline 花字 annotation — small, colored, in parentheses
        if huazi_note:
            run = p.add_run(f"（花字：{huazi_note}）")
            _apply_font(run, Pt(8), bold=False, color=COLOR_HUAZI)

    # ── Hashtags ──
    if parsed["hashtags"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(" ".join(parsed["hashtags"]))
        _apply_font(run, SIZE_HASHTAG, bold=False, color=COLOR_HASHTAG)

    doc.save(str(output_path))
    return output_path


# ═══ Parser — unchanged from v2 ═══

def _parse_llm_output(text: str, product_name: str, persona: str) -> dict:
    """Parse LLM output into structured sections.

    Expected format:
      封面：title
      简介：summary
      [body lines — one per spoken sentence, 花字 inline]
      #hashtags
    """
    lines = text.strip().split("\n")
    cover = product_name
    brief = f"{product_name}深度评测，值不值得买？"
    opening = _persona_opening(persona)
    body_lines = []
    hashtags = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("封面：") or stripped.startswith("封面:"):
            cover = stripped.replace("封面：", "").replace("封面:", "").strip()
            continue
        if stripped.startswith("简介：") or stripped.startswith("简介:"):
            brief = stripped.replace("简介：", "").replace("简介:", "").strip()
            continue

        # Hashtags — stop collecting body
        if stripped.startswith("#"):
            hashtags.append(stripped)
            continue

        # Body line — keep 花字 inline, split long lines into spoken units
        if len(stripped) > 80:
            body_lines.extend(_split_into_spoken_lines(stripped))
        else:
            body_lines.append(stripped)

    if body_lines and _is_opening(body_lines[0], persona):
        opening = body_lines.pop(0)

    return {
        "cover": cover,
        "brief": brief,
        "body": body_lines,
        "hashtags": hashtags,
    }


def _strip_prices_from_body(body_lines: list[str]) -> list[str]:
    """Move specific price mentions from voiceover to inline 花字.

    '价格呢？9999元起步' → '价格呢？比同配置便宜了好几百。（花字：9999元起）'
    """
    cleaned = []
    for line in body_lines:
        # Find price patterns: 199元, 1999元起, 3000出头
        prices = re.findall(r'\d{2,5}\s*元(?:起|起步|出头|左右)?', line)
        if prices:
            for price in prices:
                line = line.replace(price, "")
            # Clean up double punctuation and spaces
            line = re.sub(r'\s+', '', line)
            line = re.sub(r'[。！？]{2,}', lambda m: m.group(0)[0], line)
            # Append price as 花字
            price_str = " / ".join(p.strip() for p in prices)
            if line.strip():
                line = f"{line.strip()}（花字：{price_str}）"
        cleaned.append(line)
    return cleaned


def _merge_standalone_huazi(body_lines: list[str]) -> list[str]:
    """Merge standalone 花字 lines into the preceding body line.

    '核心配置，直接掀桌。' + '（花字：Core Ultra 9 / RTX 5060）'
    → '核心配置，直接掀桌。（花字：Core Ultra 9 / RTX 5060）'
    """
    if not body_lines:
        return body_lines

    merged = []
    for line in body_lines:
        stripped = line.strip()
        is_standalone_huazi = bool(re.match(r'^[（(]花字[：:].+[）)]$', stripped))

        if is_standalone_huazi and merged:
            # Append to previous line
            merged[-1] = merged[-1].rstrip() + stripped
        else:
            merged.append(line)

    return merged


def _split_into_spoken_lines(text: str) -> list[str]:
    parts = re.split(r"(?<=[。！？])", text)
    return [p.strip() for p in parts if p.strip()]


def _split_huazi_from_line(line: str) -> tuple[str, str]:
    """Split a body line into (main_text, huazi_note).

    Handles these formats:
      - "正文句子。（花字：规格）" -> ("正文句子。", "规格")
      - "正文句子" -> ("正文句子", "")
    """
    match = re.search(r'[（(]花字[：:]\s*(.+?)\s*[）)]', line)
    if match:
        huazi = match.group(1).strip()
        main = line[:match.start()].strip()
        return main, huazi
    return line, ""


def _persona_opening(persona: str) -> str:
    openings = {
        "折腾到吐": "嗨喽大家好，这里是折腾到吐。",
        "好设牛啊": "大家好，这里是好设牛啊。",
        "朋克": "哟，这里是朋克。",
        "超机懂": "大家好，我是超机懂。",
        "机能疯": "兄弟们，这里是机能疯。",
    }
    return openings.get(persona, f"嗨喽大家好，这里是{persona}。")


def _is_opening(text: str, persona: str) -> bool:
    patterns = [
        "嗨喽大家好", "大家好", "这里是折腾到吐",
        "这里是好设牛啊", "这里是朋克", "我是超机懂",
        "这里是机能疯",
    ]
    return any(p in text for p in patterns)
