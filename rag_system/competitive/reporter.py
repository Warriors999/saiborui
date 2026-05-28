"""Generate professional competitive analysis reports in .docx format.

Swiss Editorial design system: 等线 font, #1E40AF deep blue accent,
color-coded callout boxes, styled quotes, clean data tables.
"""

import json
from collections import Counter
from datetime import datetime, timedelta
from pathlib import Path

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from rag_system.competitive.models import WeeklyReport
from rag_system.competitive.store import get_analyzed_videos, save_report

REPORTS_DOCX_DIR = Path("output/competitive/reports")

# ── Design System ──
PRIMARY = "1E40AF"
PRIMARY_LIGHT = "DBEAFE"
ACCENT = "3B82F6"
BG_WHITE = "FFFFFF"
BG_SURFACE = "F8FAFC"
TEXT_PRIMARY = "1E293B"
TEXT_SECONDARY = "64748B"
QUOTE_BG = "F1F5F9"
CALLOUT_KEY_BG = "1E40AF"
CALLOUT_VAL_BG = "EFF6FF"
TABLE_HEADER_BG = "1E40AF"
TABLE_ROW_ALT = "F8FAFC"


def generate_docx_report(videos: list[dict] = None, period: str = "weekly") -> Path:
    """Generate a professionally formatted .docx competitive analysis report."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    if videos is None:
        videos = get_analyzed_videos()

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = '等线'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(30, 41, 59)
    style.paragraph_format.space_after = Pt(6)

    # ── Helper: colored paragraph with background ──
    def add_callout(doc, title_text, body_text, bg_color=CALLOUT_VAL_BG, title_color=PRIMARY):
        """Add a visually distinct callout box."""
        # Title bar
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(0)
        _shade_paragraph(p_title, CALLOUT_KEY_BG)
        run = p_title.add_run(f"  {title_text}")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = '等线'

        # Body
        p_body = doc.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(8)
        _shade_paragraph(p_body, bg_color)
        run = p_body.add_run(body_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(30, 41, 59)
        run.font.name = '等线'

    def add_quote(doc, text):
        """Add a styled quote (italic, gray, indented)."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _shade_paragraph(p, QUOTE_BG)
        # Vertical bar via left border
        pPr = p._element.get_or_add_pPr()
        from lxml import etree
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
        left = etree.SubElement(pBdr, qn('w:left'))
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), PRIMARY)
        run = p.add_run(f"  {text}")
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)
        run.font.name = '等线'

    def add_table(doc, headers, rows):
        """Add a clean styled table."""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = '等线'
            _shade_cell(cell, TABLE_HEADER_BG)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for ri, row in enumerate(rows):
            bg = TABLE_ROW_ALT if ri % 2 == 0 else BG_WHITE
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(9)
                run.font.name = '等线'
                _shade_cell(cell, bg)

        doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('竞品学习报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)
    run.font.name = '等线'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Competitive Learning Report · {period.upper()}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime('%Y.%m.%d'))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # OVERVIEW
    # ═══════════════════════════════════════════
    doc.add_heading('概览', level=1)
    doc.add_paragraph(f'本报告分析了 {len(videos)} 个竞品视频，覆盖文案脚本、视觉节奏、剪辑手法。')

    # Video list table
    if videos:
        headers = ['视频', '创作者', '播放量', '钩子类型', '口语密度']
        rows = []
        for v in videos:
            rows.append([
                v.get('title', '')[:50],
                v.get('creator', ''),
                f"{v.get('views', 0):,}",
                v.get('hook_type', ''),
                f"{v.get('spoken_density', 0):.1f}",
            ])
        add_table(doc, headers, rows)

    # ═══════════════════════════════════════════
    # VIDEO DETAILS
    # ═══════════════════════════════════════════
    for vi, v in enumerate(videos, 1):
        title_text = v.get('title', '')[:80]
        creator = v.get('creator', '')
        views = v.get('views', 0)
        hook = v.get('hook_type', '')
        deep = v.get('deep_analysis', '')
        visual_raw = v.get('visual_analysis', '')

        doc.add_heading(f'{vi}. {title_text}', level=2)

        # Creator meta
        meta = doc.add_paragraph()
        run = meta.add_run(f'{creator}    |    {views:,} 播放    |    钩子: {hook}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)

        # Visual data table
        if visual_raw:
            try:
                vis = json.loads(visual_raw) if isinstance(visual_raw, str) else visual_raw
                if not vis.get('error'):
                    add_table(doc,
                        ['指标', '数值'],
                        [
                            ['镜头数', str(vis.get('shot_count', '-'))],
                            ['平均镜头时长', f"{vis.get('avg_shot_sec', '-')}s"],
                            ['剪辑频率', f"{vis.get('cuts_per_minute', '-')} 次/分钟"],
                            ['短镜头(≤2s)占比', f"{vis.get('short_shots_pct', '-')}%"],
                            ['中镜头(2-5s)占比', f"{vis.get('medium_shots_pct', '-')}%"],
                            ['长镜头(>5s)占比', f"{vis.get('long_shots_pct', '-')}%"],
                        ])
            except (json.JSONDecodeError, TypeError):
                pass

        # Deep analysis — parse into sections and render as callouts
        if deep:
            sections = _parse_deep_sections(deep)
            for sec_title, sec_body in sections:
                if '改编' in sec_title or '学习' in sec_title:
                    add_callout(doc, sec_title, sec_body, bg_color="FEF3C7")
                elif '句式' in sec_title or '技巧' in sec_title:
                    add_callout(doc, sec_title, sec_body, bg_color=CALLOUT_VAL_BG)
                elif '视觉' in sec_title or '剪辑' in sec_title or '拍摄' in sec_title:
                    add_callout(doc, sec_title, sec_body, bg_color="F0FDF4")
                elif '钩子' in sec_title:
                    add_callout(doc, sec_title, sec_body, bg_color="FEF2F2")
                else:
                    add_callout(doc, sec_title, sec_body)

        # Extract and render quotes
        if deep:
            for line in deep.split('\n'):
                if '原文：' in line or '原文:' in line:
                    quote = line.split('：', 1)[-1].split(':', 1)[-1].strip().strip('"').strip('"')
                    if len(quote) > 15:
                        add_quote(doc, quote)

        doc.add_paragraph()

    # ═══════════════════════════════════════════
    # RECOMMENDATIONS
    # ═══════════════════════════════════════════
    doc.add_heading('总结与建议', level=1)
    report = generate_weekly_report()
    for rec in report.recommendations:
        p = doc.add_paragraph()
        run = p.add_run(f'▸ {rec}')
        run.font.size = Pt(10)

    # ═══════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 赛博瑞 数据驱动内容工厂 · 竞品学习系统 —')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    # Save
    REPORTS_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d")
    path = REPORTS_DOCX_DIR / f"竞品学习报告_{period}_{date_str}.docx"
    doc.save(str(path))
    return path


def _shade_paragraph(paragraph, color: str):
    """Add background shading to a paragraph."""
    from lxml import etree
    pPr = paragraph._element.get_or_add_pPr()
    shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')


def _shade_cell(cell, color: str):
    """Add background shading to a table cell."""
    from lxml import etree
    tcPr = cell._element.get_or_add_tcPr()
    shd = etree.SubElement(tcPr, qn('w:shd'))
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')


def _parse_deep_sections(deep_text: str) -> list[tuple[str, str]]:
    """Parse LLM deep analysis into titled sections."""
    sections = []
    current_title = "学习要点"
    current_body = []

    for line in deep_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Detect section headers like "1. 脚本结构拆解"
        if line[0].isdigit() and '. ' in line[:4]:
            if current_body:
                sections.append((current_title, '\n'.join(current_body)))
            current_title = line.split('. ', 1)[1] if '. ' in line else line
            current_body = []
        else:
            cleaned = line.lstrip('- ').strip()
            if cleaned:
                current_body.append(cleaned)

    if current_body:
        sections.append((current_title, '\n'.join(current_body)))

    return sections if sections else [("完整分析", deep_text)]


def generate_weekly_report() -> WeeklyReport:
    """Generate a weekly summary report."""
    now = datetime.now()
    week_ago = (now - timedelta(days=7)).isoformat()
    all_videos = get_analyzed_videos()
    recent = [v for v in all_videos if v.get("analyzed_at", "") >= week_ago]

    report = WeeklyReport(
        period_start=week_ago[:10],
        period_end=now.isoformat()[:10],
        videos_analyzed=len(recent),
    )

    if not recent:
        report.recommendations = ["本周无新分析的视频。运行 search 命令开始收集。"]
        save_report(report)
        return report

    creator_counts = Counter(v.get("creator", "unknown") for v in recent)
    report.top_creators = [name for name, _ in creator_counts.most_common(5)]

    hook_counts = Counter(v.get("hook_type", "unknown") for v in recent)
    report.trending_hook_types = dict(hook_counts.most_common())

    cat_groups = {}
    for v in recent:
        cat = v.get("category", "other")
        if cat not in cat_groups:
            cat_groups[cat] = []
        cat_groups[cat].append(v)
    for cat, vids in cat_groups.items():
        avg_spoken = sum(v.get("spoken_density", 0) for v in vids) / len(vids)
        top_hook = Counter(v.get("hook_type", "") for v in vids).most_common(1)[0][0]
        report.category_insights[cat] = {
            "count": len(vids),
            "avg_spoken_density": round(avg_spoken, 2),
            "dominant_hook": top_hook,
        }

    all_patterns = []
    for v in recent:
        all_patterns.extend(v.get("standout_patterns", []))
    pattern_counts = Counter(all_patterns)
    report.new_patterns_discovered = [p for p, _ in pattern_counts.most_common(5)]

    report.recommendations = _generate_recommendations(report)
    save_report(report)
    return report


def _generate_recommendations(report: WeeklyReport) -> list[str]:
    recs = []
    if report.trending_hook_types:
        top_hook = max(report.trending_hook_types, key=report.trending_hook_types.get)
        recs.append(f"本周最热门钩子类型是「{top_hook}」，建议下期脚本优先尝试该开篇方式")
    analyzed_cats = set(report.category_insights.keys())
    all_cats = {"keyboard", "mouse", "monitor", "laptop", "phone", "gpu", "headphone", "desk_chair"}
    missing = all_cats - analyzed_cats
    if missing:
        recs.append(f"以下品类本周无竞品数据: {', '.join(missing)}，建议补充搜索")
    if report.top_creators:
        recs.append(f"值得关注的创作者: {', '.join(report.top_creators[:3])}")
    return recs
